#!/usr/bin/env python3
"""LQA Inspector - Glossary Parser.

Deterministic parser for glossary source files (XLSX, CSV, DOCX, PDF, TXT).
Two-phase flow for structured sources:
  1. parse_headers: detect columns and show sample data
  2. import_with_mapping: extract entries using user-provided column mapping
Direct extraction for unstructured sources (DOCX, PDF, TXT).
"""

import json
import os
import sys
import csv
from typing import Optional

# Ensure UTF-8 stdout on Windows (cp1252 cannot encode CJK/emoji/Indonesian chars)
if hasattr(sys.stdout, 'reconfigure'):
    sys.stdout.reconfigure(encoding='utf-8')


def detect_source_type(filepath: str) -> str:
    ext = os.path.splitext(filepath)[1].lower()
    return {
        ".xlsx": "xlsx", ".xls": "xlsx", ".csv": "csv",
        ".docx": "docx", ".pdf": "pdf", ".txt": "txt",
    }.get(ext, "unknown")


def _cell_str(row: list, col: int) -> str:
    if col < 0 or col >= len(row):
        return ""
    val = row[col]
    return str(val).strip() if val is not None else ""


def _read_xlsx_rows(filepath: str, sheet_name: Optional[str] = None) -> list:
    import openpyxl
    wb = openpyxl.load_workbook(filepath, data_only=True)
    sheets = wb.sheetnames
    target = sheet_name if sheet_name and sheet_name in sheets else sheets[0]
    ws = wb[target]
    rows = []
    for row in ws.iter_rows(values_only=True):
        rows.append([c for c in row])
    wb.close()
    return rows


def _read_csv_rows(filepath: str) -> list:
    encoding = "utf-8"
    try:
        with open(filepath, "r", encoding="utf-8") as f:
            f.read(1024)
    except UnicodeDecodeError:
        encoding = "latin-1"
    rows = []
    with open(filepath, "r", encoding=encoding, newline="") as f:
        sample = f.read(8192)
        f.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        except csv.Error:
            dialect = csv.excel
        reader = csv.reader(f, dialect)
        for row in reader:
            rows.append(row)
    return rows


def _detect_encoding(filepath: str) -> str:
    for enc in ("utf-8", "utf-8-sig", "latin-1"):
        try:
            with open(filepath, "r", encoding=enc) as f:
                f.read(1024)
            return enc
        except UnicodeDecodeError:
            continue
    return "latin-1"


def parse_headers(filepath: str, sheet_name: Optional[str] = None) -> dict:
    if not os.path.isfile(filepath):
        return {"success": False, "error": f"File not found: {filepath}"}
    source_type = detect_source_type(filepath)
    filename = os.path.basename(filepath)
    try:
        if source_type == "xlsx":
            return _parse_xlsx_headers(filepath, filename, sheet_name)
        elif source_type == "csv":
            return _parse_csv_headers(filepath, filename)
        else:
            return {"success": True, "filename": filename, "source_type": source_type,
                    "headers": [], "sample_rows": [], "sheets": [], "requires_mapping": False}
    except Exception as e:
        return {"success": False, "error": f"Failed to parse headers: {e}"}


def _parse_xlsx_headers(filepath: str, filename: str, sheet_name: Optional[str]) -> dict:
    import openpyxl
    wb = openpyxl.load_workbook(filepath, read_only=True, data_only=True)
    sheets = wb.sheetnames
    target = sheet_name if sheet_name else sheets[0]
    if target not in sheets:
        wb.close()
        return {"success": False, "error": f"Sheet '{target}' not found. Available: {sheets}"}
    ws = wb[target]
    headers, sample_rows = [], []
    for row_idx, row in enumerate(ws.iter_rows(values_only=True), 1):
        if row_idx == 1:
            headers = [str(c) if c is not None else f"Col{i+1}" for i, c in enumerate(row)]
        elif row_idx <= 6:
            sample_rows.append([str(c) if c is not None else "" for c in row])
        if row_idx > 6:
            break
    wb.close()
    return {"success": True, "filename": filename, "source_type": "xlsx",
            "headers": headers, "sample_rows": sample_rows,
            "sheets": sheets, "active_sheet": target, "requires_mapping": True}


def _parse_csv_headers(filepath: str, filename: str) -> dict:
    encoding = _detect_encoding(filepath)
    with open(filepath, "r", encoding=encoding, newline="") as f:
        sample = f.read(8192)
        f.seek(0)
        try:
            dialect = csv.Sniffer().sniff(sample, delimiters=",;\t|")
        except csv.Error:
            dialect = csv.excel
        reader = csv.reader(f, dialect)
        headers, sample_rows = [], []
        for i, row in enumerate(reader):
            if i == 0:
                headers = [c.strip() for c in row]
            elif i <= 5:
                sample_rows.append([c.strip() for c in row])
            if i > 5:
                break
    return {"success": True, "filename": filename, "source_type": "csv",
            "headers": headers, "sample_rows": sample_rows, "sheets": [], "requires_mapping": True}


def import_with_mapping(filepath: str, column_mapping: dict, source_type: str = None, sheet_name: str = None) -> dict:
    """Phase 2: Extract glossary entries using user-provided column mapping."""
    if not os.path.isfile(filepath):
        return {"success": False, "error": f"File not found: {filepath}", "entries": [], "warnings": []}
    if source_type is None:
        source_type = detect_source_type(filepath)
    filename = os.path.basename(filepath)
    warnings = []
    source_col = column_mapping.get("source_term_col")
    if source_col is None:
        return {"success": False, "error": "source_term_col required", "entries": [], "warnings": []}
    target_col = column_mapping.get("target_term_col")
    status_col = column_mapping.get("status_col")
    dnt_col = column_mapping.get("dnt_col")
    banned_col = column_mapping.get("banned_col")
    notes_col = column_mapping.get("notes_col")
    try:
        if source_type == "xlsx":
            rows = _read_xlsx_rows(filepath, sheet_name)
        elif source_type == "csv":
            rows = _read_csv_rows(filepath)
        else:
            return {"success": False, "error": f"Column mapping not for {source_type}", "entries": [], "warnings": []}
        if not rows:
            return {"success": False, "error": "No data rows", "entries": [], "warnings": []}
        entries = []
        for row_idx, row in enumerate(rows[1:], 2):
            if source_col >= len(row):
                warnings.append(f"Row {row_idx}: col {source_col} out of range")
                continue
            source_term = _cell_str(row, source_col)
            if not source_term:
                continue
            target_term = _cell_str(row, target_col) if target_col is not None and target_col < len(row) else ""
            status_str = _cell_str(row, status_col).lower() if status_col is not None and status_col < len(row) else "approved"
            dnt_raw = _cell_str(row, dnt_col).lower() if dnt_col is not None and dnt_col < len(row) else ""
            banned_raw = _cell_str(row, banned_col).lower() if banned_col is not None and banned_col < len(row) else ""
            notes = _cell_str(row, notes_col) if notes_col is not None and notes_col < len(row) else ""
            dnt = status_str in ("dnt", "do_not_translate") or dnt_raw in ("yes", "true", "1", "x", "dnt")
            ban = status_str in ("banned", "forbidden") or banned_raw in ("yes", "true", "1", "x", "banned")
            if dnt:
                status = "do_not_translate"
            elif ban:
                status = "banned"
            elif status_str in ("provisional", "pending", "review"):
                status = "provisional"
            elif status_str in ("deprecated", "obsolete"):
                status = "deprecated"
            else:
                status = "approved"
            entries.append({"source_term": source_term, "target_term": target_term, "status": status,
                           "do_not_translate": dnt, "banned": ban, "notes": notes, "source_row": row_idx})
        return {"success": True, "entries": entries, "warnings": warnings,
                "metadata": {"filename": filename, "source_type": source_type, "total_entries": len(entries)}}
    except Exception as e:
        return {"success": False, "error": f"Import failed: {e}", "entries": [], "warnings": warnings}


def extract_text_content(filepath: str, source_type: str = None) -> dict:
    """Extract text from unstructured sources (DOCX, PDF, TXT)."""
    if not os.path.isfile(filepath):
        return {"success": False, "error": f"File not found: {filepath}", "chunks": [], "warnings": []}
    if source_type is None:
        source_type = detect_source_type(filepath)
    filename = os.path.basename(filepath)
    warnings = []
    try:
        if source_type == "docx":
            return _extract_docx(filepath, filename, warnings)
        elif source_type == "pdf":
            return _extract_pdf(filepath, filename, warnings)
        elif source_type == "txt":
            return _extract_txt(filepath, filename, warnings)
        else:
            return {"success": False, "error": f"Unsupported: {source_type}", "chunks": [], "warnings": warnings}
    except Exception as e:
        return {"success": False, "error": f"Extraction failed: {e}", "chunks": [], "warnings": warnings}


def _extract_docx(filepath: str, filename: str, warnings: list) -> dict:
    from docx import Document as DocxDocument
    doc = DocxDocument(filepath)
    chunks, order = [], 0
    heading, text_parts = "", []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = para.style.name if para.style else ""
        if "Heading" in style or style.startswith("Title"):
            if text_parts:
                chunks.append({"source_page": 0, "source_sheet": "", "source_row_start": 0,
                               "source_row_end": 0, "heading": heading,
                               "content": "\n".join(text_parts), "chunk_order": order})
                order += 1
                text_parts = []
            heading = text
        else:
            text_parts.append(text)
    if text_parts:
        chunks.append({"source_page": 0, "source_sheet": "", "source_row_start": 0,
                       "source_row_end": 0, "heading": heading,
                       "content": "\n".join(text_parts), "chunk_order": order})
        order += 1
    for ti, table in enumerate(doc.tables):
        rows_text = ["\t".join(cell.text.strip() for cell in row.cells) for row in table.rows]
        if rows_text:
            chunks.append({"source_page": 0, "source_sheet": "", "source_row_start": 0,
                           "source_row_end": len(rows_text), "heading": f"Table {ti+1}",
                           "content": "\n".join(rows_text), "chunk_order": order})
            order += 1
    if not chunks:
        warnings.append("No text content extracted from DOCX file.")
    return {"success": True, "chunks": chunks, "warnings": warnings,
            "metadata": {"filename": filename, "source_type": "docx", "total_chunks": len(chunks)}}


def _extract_pdf(filepath: str, filename: str, warnings: list) -> dict:
    from pypdf import PdfReader
    reader = PdfReader(filepath)
    chunks, order = [], 0
    for page_num, page in enumerate(reader.pages, 1):
        text = page.extract_text()
        if text and text.strip():
            paras = [p.strip() for p in text.split("\n\n") if p.strip()]
            if not paras:
                paras = [text.strip()]
            for para in paras:
                chunks.append({"source_page": page_num, "source_sheet": "", "source_row_start": 0,
                               "source_row_end": 0, "heading": f"Page {page_num}",
                               "content": para, "chunk_order": order})
                order += 1
        else:
            warnings.append(f"Page {page_num}: no text extracted.")
    if not chunks:
        warnings.append("No text content extracted from PDF.")
    return {"success": True, "chunks": chunks, "warnings": warnings,
            "metadata": {"filename": filename, "source_type": "pdf", "total_chunks": len(chunks)}}


def _extract_txt(filepath: str, filename: str, warnings: list) -> dict:
    encoding = _detect_encoding(filepath)
    with open(filepath, "r", encoding=encoding) as f:
        content = f.read()
    paras = [p.strip() for p in content.split("\n\n") if p.strip()]
    if not paras and content.strip():
        paras = [content.strip()]
    chunks = [{"source_page": 0, "source_sheet": "", "source_row_start": 0,
               "source_row_end": 0, "heading": "", "content": p, "chunk_order": i}
              for i, p in enumerate(paras)]
    if not chunks:
        warnings.append("No text content extracted from TXT file.")
    return {"success": True, "chunks": chunks, "warnings": warnings,
            "metadata": {"filename": filename, "source_type": "txt", "total_chunks": len(chunks)}}


def main():
    if len(sys.argv) < 3:
        print(json.dumps({"success": False, "error": "Usage: glossary_parser.py <command> <filepath> [args...]"}))
        sys.exit(1)
    command, filepath = sys.argv[1], sys.argv[2]
    if command == "parse_headers":
        sheet = sys.argv[3] if len(sys.argv) > 3 else None
        result = parse_headers(filepath, sheet)
    elif command == "import_with_mapping":
        mapping = json.loads(sys.argv[3]) if len(sys.argv) > 3 else {}
        sheet = sys.argv[4] if len(sys.argv) > 4 else None
        result = import_with_mapping(filepath, mapping, sheet_name=sheet)
    elif command == "extract_text":
        result = extract_text_content(filepath)
    else:
        result = {"success": False, "error": f"Unknown command: {command}"}
    print(json.dumps(result, ensure_ascii=False))


if __name__ == "__main__":
    main()